from docx import Document
from pdf2docx import parse
from docx2pdf import convert
from tqdm import tqdm
import re
import torch

try:
    from .model import translate
except:
    #To enable relative imports on Colab
    from model import translate

#language to model recognized type
support_langs = {'arabic': 'ar_AR', 'czech': 'cs_CZ', 'german': 'de_DE', 'english': 'en_XX', 'spanish': 'es_XX', 'estonian': 'et_EE', 
                 'finnish': 'fi_FI', 'french': 'fr_XX', 'gujarati': 'gu_IN', 'hindi': 'hi_IN', 'italian': 'it_IT', 'japanese': 'ja_XX', 
                 'kazakh': 'kk_KZ', 'korean': 'ko_KR', 'lithuanian': 'lt_LT', 'latvian': 'lv_LV', 'burmese': 'my_MM', 'nepali': 'ne_NP', 
                 'dutch': 'nl_XX', 'romanian': 'ro_RO', 'russian': 'ru_RU', 'sinhala': 'si_LK', 'turkish': 'tr_TR', 'vietnamese': 'vi_VN', 
                 'chinese': 'zh_CN', 'afrikaans': 'af_ZA', 'azerbaijani': 'az_AZ', 'bengali': 'bn_IN', 'persian': 'fa_IR', 'hebrew': 'he_IL', 
                 'croatian': 'hr_HR', 'indonesian': 'id_ID', 'georgian': 'ka_GE', 'khmer': 'km_KH', 'macedonian': 'mk_MK', 'malayalam': 'ml_IN', 
                 'mongolian': 'mn_MN', 'marathi': 'mr_IN', 'polish': 'pl_PL', 'pashto': 'ps_AF', 'portuguese': 'pt_XX', 'swedish': 'sv_SE', 
                 'swahili': 'sw_KE', 'tamil': 'ta_IN', 'telugu': 'te_IN', 'thai': 'th_TH', 'tagalog': 'tl_XX', 'ukrainian': 'uk_UA', 
                 'urdu': 'ur_PK', 'xhosa': 'xh_ZA', 'galician': 'gl_ES', 'slovene': 'sl_SI'}


def pdf_to_docx(fp):
    try:
        print('Converting pdf to docx...')
        docx_fp = fp.replace('pdf', 'docx')
        parse(fp, docx_fp)
        return docx_fp
    except:
        return ''
    

def get_docx(fp):    
    if fp[-3:] == 'pdf':
        fp = pdf_to_docx(fp)
    elif fp[-4:] != 'docx':
        raise Exception('Unsupported document')
    
    #parse docx file as docx Document object
    return Document(fp)


def translate_doc(fp, in_lang, out_lang, output_fp='', batch=10):   
    """Translate pdf or docx documents

    Args:
        fp (str): source file path
        in_lang (str): source languge
        out_lang (str): targte language
        output_fp (str, optional): target file path. Defaults to same location as fp, file name default append target_language.
        batch (int, optional): batch of sequences for parallel processing, set smaller if CUDA out of memory. if run on cpu, set batch to 1.
    
    """     
    #allow upper/lower case
    if in_lang.lower() not in support_langs:
        raise Exception(f'{in_lang} language is not supported.')
    if out_lang.lower() not in support_langs:
        raise Exception(f'{out_lang} language is not supported.')
    
    in_lang = support_langs[in_lang]
    out_lang = support_langs[out_lang]
    doc = get_docx(fp)    
    device = "cuda" if torch.cuda.is_available() else "cpu"
    if device == 'cpu': #does not support parallel
        batch = 1

    print('Translating paragraphs in parallel....')            
    idx = []
    in_paras = []
    #gather paragraph content as input sequence 
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            idx.append(i)
            in_paras.append(p.text.strip())
    
    out_paras = []
    #translate paragraphs by batch
    for i in tqdm(range(0, len(idx), batch)):
        out_paras += translate(in_paras[i: i+batch], in_lang, out_lang, device=device)                       
    for i, id in enumerate(idx):
        doc.paragraphs[id].text = out_paras[i]

    
    print('Translating table by table....')
    locs = []
    in_vals = []    
    for i, t in tqdm(enumerate(doc.tables)):  
        print('Table ', i+1)      
        m = len(t.rows)
        n = len(t.columns)
        locs = []
        in_vals = []        
        for j in range(m):
            for k in range(n):
                val = t.cell(j,k).text     
                #if cell values not numbers or special chars, translate       
                if re.findall(r'[^0-9,.%$:()\n -]+', val):                    
                    locs.append((j,k))
                    in_vals.append(val.strip())
        
        #translate cell values by batch
        out_vals = []
        for b in tqdm(range(0, len(locs), batch)):
            out_vals += translate(in_vals[b:b+batch], in_lang, out_lang, device=device)
        
        for i, (j,k) in enumerate(locs):                                 
            #only replace text, try to preserve original tab or spaces   
            t.cell(j,k).text = t.cell(j,k).text.replace(in_vals[i], out_vals[i])

    if not output_fp:
        output_fp = f"{fp.split('.')[0]}_{out_lang}.docx"
    doc.save(output_fp)
    if fp[-3:] == 'pdf':
        try:            
            convert(output_fp, output_fp.replace('docx', 'pdf'))
            output_fp = output_fp.replace('docx', 'pdf')
        except Exception as e:
            #docx2pdf requires local installation of word
            print(e)

    msg = f'Translation finished, please find the translated file at "{output_fp}".'
    print(msg)
    return msg


if __name__ == '__main__':
    translate_doc('samples/apple.pdf', 'english', 'chinese')